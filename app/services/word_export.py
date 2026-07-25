# -*- coding: utf-8 -*-
"""Экспорт протокола разногласий и сопроводительного письма в один Word документ."""
import re
import urllib.parse
import logging
from io import BytesIO
from fastapi import HTTPException
from fastapi.responses import Response
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

FONT_NAME = "Times New Roman"
FONT_SIZE = 12


def _clean_text(text: str) -> str:
    """Очистка текста от множественных пробелов, markdown и спецсимволов."""
    # Убираем markdown форматирование
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # **жирный**
    text = re.sub(r'\*([^*]+)\*', r'\1', text)  # *курсив*
    text = re.sub(r'^---+$', '', text)  # ---
    text = re.sub(r'^#+\s*', '', text)  # # заголовки
    # Убираем все виды множественных пробелов (обычные, неразрывные и т.д.)
    text = re.sub(r'\s+', ' ', text)
    # Убираем пробелы в начале и конце
    text = text.strip()
    return text


def _add_paragraph(doc: Document, text: str, space_after: int = 0, alignment: int = 0):
    text = _clean_text(text)
    if not text:
        return
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.0  # Одинарный межстроковый интервал


def _make_docx_response(buffer: BytesIO, filename: str) -> Response:
    buffer.seek(0)
    encoded = urllib.parse.quote(filename, safe="")
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )


def _add_letter_header(doc: Document, letter_text: str):
    """Добавить шапку письма с правильным форматированием."""
    lines = letter_text.split("\n")
    header_lines = []
    content_start = 0
    
    # Убираем лишние заголовки
    filtered_lines = []
    for line in lines:
        line_clean = _clean_text(line.upper())
        # Пропускаем строки с заголовками письма и датой в шапке
        if ("СОПРОВОДИТЕЛЬНОЕ ПИСЬМО" in line_clean or 
            "К ПРОТОКОЛУ РАЗНОГЛАСИЙ" in line_clean or
            ("___" in line and "202" in line and len(line_clean) < 50)):  # Дата в шапке
            continue
        filtered_lines.append(line)
    
    lines = filtered_lines
    
    # Находим шапку (до "Уважаемые коллеги" или аналогичной строки)
    for i, line in enumerate(lines):
        if "уважаем" in line.lower() or "добрый день" in line.lower() or "доброго" in line.lower():
            content_start = i
            break
        if line.strip():
            header_lines.append(line.strip())
    
    # Добавляем шапку с выравниванием по правому краю (без межстроковых интервалов)
    for i, line in enumerate(header_lines):
        line = _clean_text(line)
        if line:
            _add_paragraph(doc, line, space_after=0, alignment=2)
    
    # Отступ после шапки
    if header_lines:
        # Изменяем последний параграф шапки - добавляем отступ
        last_para = doc.paragraphs[-1]
        last_para.paragraph_format.space_after = Pt(12)
    
    # Добавляем приветствие по центру
    if content_start < len(lines):
        greeting_line = _clean_text(lines[content_start])
        if greeting_line:
            _add_paragraph(doc, greeting_line, space_after=12, alignment=1)
    
    # Возвращаем оставшийся текст
    return "\n".join(lines[content_start + 1:])


def export_protocol_and_letter(
    protocol_text: str,
    letter_text: str,
    filename: str,
    part: str = "both",
) -> Response:
    """
    Создать .docx: part='both' — оба раздела, 'protocol' — только протокол, 'letter' — только письмо.
    """
    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE)
        # Сбрасываем все интервалы в стиле Normal
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

        if part in ("both", "protocol"):
            # Добавляем заголовок
            para = doc.add_paragraph()
            run = para.add_run("Протокол разногласий")
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
            run.font.bold = False
            para.alignment = 0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            
            # Пустая строка после заголовка
            doc.add_paragraph()
            
            # Обрабатываем протокол
            lines = (protocol_text or "").split("\n")
            
            for line in lines:
                line = _clean_text(line)
                if not line:
                    continue
                
                line_upper = line.upper()
                
                # Перед "ПУНКТ ДОГОВОРА" добавляем пустую строку
                if "ПУНКТ ДОГОВОРА" in line_upper and len(doc.paragraphs) > 2:
                    doc.add_paragraph()
                
                # Перед "Редакция Поставщика" или "Редакция Покупателя" добавляем пустую строку
                if ("РЕДАКЦИЯ ПОСТАВЩИКА" in line_upper or "РЕДАКЦИЯ ПОКУПАТЕЛЯ" in line_upper) and len(doc.paragraphs) > 2:
                    doc.add_paragraph()
                
                # Добавляем строку
                _add_paragraph(doc, line, space_after=0)


        if part == "both":
            # Разрыв страницы между протоколом и письмом
            doc.add_page_break()

        if part in ("both", "letter"):
            if letter_text:
                remaining_text = _add_letter_header(doc, letter_text)
                
                # Ищем "С уважением" или "ЗАКЛЮЧЕНИЕ" для разделения
                lines = remaining_text.split("\n")
                conclusion_start = -1
                signature_start = -1
                
                for i, line in enumerate(lines):
                    line_clean = _clean_text(line.upper())
                    if "ЗАКЛЮЧЕНИЕ" in line_clean and conclusion_start == -1:
                        conclusion_start = i
                    if "С УВАЖЕНИЕМ" in line_clean and signature_start == -1:
                        signature_start = i
                        break
                
                # Обрабатываем основной текст (до заключения или подписи)
                end_main = len(lines)
                if signature_start != -1:
                    end_main = signature_start
                elif conclusion_start != -1:
                    end_main = conclusion_start
                
                for i in range(end_main):
                    line = _clean_text(lines[i])
                    
                    if not line:
                        # Пропускаем пустые строки
                        continue
                    
                    line_lower = line.lower()
                    
                    # Если строка начинается с цифры (пункт аргументации), добавляем больший отступ после предыдущего
                    if re.match(r'^\d+\.', line) and i > 0 and len(doc.paragraphs) > 0:
                        doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
                    
                    # Перед финальными абзацами добавляем пустую строку
                    if ("просим рассмотреть" in line_lower or 
                        "в случае возникновения" in line_lower or
                        "готовы к конструктивному" in line_lower) and len(doc.paragraphs) > 0:
                        doc.add_paragraph()
                    
                    _add_paragraph(doc, line, space_after=3)
                
                # Обрабатываем заключение и подпись
                if signature_start != -1:
                    # Добавляем отступ перед подписью
                    if len(doc.paragraphs) > 0:
                        doc.paragraphs[-1].paragraph_format.space_after = Pt(18)
                    
                    # Собираем строки подписи
                    signature_lines = []
                    for i in range(signature_start, len(lines)):
                        line = _clean_text(lines[i])
                        if line:
                            signature_lines.append(line)
                    
                    # Ищем должность и ФИО для подписи
                    position_lines = []
                    signature_line = ""
                    remaining_lines = []
                    
                    found_position = False
                    for line in signature_lines:
                        line_upper = line.upper()
                        if not found_position and ("С УВАЖЕНИЕМ" in line_upper or 
                                                   "УПРАВЛЯЮЩ" in line_upper or 
                                                   "ДИРЕКТОР" in line_upper or
                                                   "ПРЕДПРИНИМАТЕЛЬ" in line_upper or
                                                   "ФОНД" in line_upper):
                            position_lines.append(line)
                            if "ФОНД" in line_upper or "ООО" in line_upper or "ИП" in line_upper:
                                found_position = True
                        elif found_position and "/" in line and "_" in line:
                            signature_line = line
                        elif found_position:
                            remaining_lines.append(line)
                    
                    # Добавляем должность слева и подпись справа в одной таблице
                    if signature_line:
                        signature_line = _clean_text(signature_line)
                        # Создаем таблицу для выравнивания
                        table = doc.add_table(rows=1, cols=2)
                        table.autofit = False
                        table.allow_autofit = False
                        
                        # Убираем границы таблицы
                        for row in table.rows:
                            for cell in row.cells:
                                tc = cell._element
                                tcPr = tc.get_or_add_tcPr()
                                tcBorders = OxmlElement('w:tcBorders')
                                for border_name in ['top', 'left', 'bottom', 'right']:
                                    border = OxmlElement(f'w:{border_name}')
                                    border.set(qn('w:val'), 'none')
                                    tcBorders.append(border)
                                tcPr.append(tcBorders)
                        
                        # Левая ячейка - должность (все строки)
                        left_cell = table.rows[0].cells[0]
                        left_cell.width = Inches(3.5)
                        # Удаляем стандартный параграф
                        left_cell._element.clear_content()
                        for idx, line in enumerate(position_lines):
                            line = _clean_text(line)
                            para = left_cell.add_paragraph()
                            run = para.add_run(line)
                            run.font.name = FONT_NAME
                            run.font.size = Pt(FONT_SIZE)
                            para.paragraph_format.space_before = Pt(0)
                            para.paragraph_format.space_after = Pt(0)
                            para.paragraph_format.line_spacing = 1.0
                        
                        # Правая ячейка - линия подписи и ФИО (выше на уровне последней строки должности)
                        right_cell = table.rows[0].cells[1]
                        right_cell.width = Inches(2.5)
                        para = right_cell.paragraphs[0]
                        run = para.add_run(signature_line)
                        run.font.name = FONT_NAME
                        run.font.size = Pt(FONT_SIZE)
                        para.alignment = 2  # По правому краю
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.line_spacing = 1.0
                        # Выравниваем по нижнему краю ячейки
                        tc = right_cell._element
                        tcPr = tc.get_or_add_tcPr()
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), 'bottom')
                        tcPr.append(vAlign)
                    
                    # Добавляем отступ перед М.П.
                    if len(doc.paragraphs) > 0:
                        doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
                    
                    # Добавляем М.П. и дату (слева внизу)
                    for line in remaining_lines:
                        line = _clean_text(line)
                        _add_paragraph(doc, line, space_after=0, alignment=0)

        buffer = BytesIO()
        doc.save(buffer)

        if part == "protocol":
            doc_filename = f"{filename}_протокол.docx"
        elif part == "letter":
            doc_filename = f"{filename}_письмо.docx"
        else:
            doc_filename = f"{filename}.docx"

        return _make_docx_response(buffer, doc_filename)
    except ImportError:
        raise HTTPException(status_code=500, detail="Модуль python-docx не установлен")
    except Exception as e:
        logger.exception("Word export error")
        raise HTTPException(status_code=500, detail=f"Ошибка создания Word: {str(e)}")
